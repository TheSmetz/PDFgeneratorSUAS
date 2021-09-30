from flask import Flask, render_template, send_file, redirect, request
from docx import Document
import time

# TODO: MODIFY ALL THE DOCUMENTS ADDING THE DATA (DEAR NAME SURNAME)

app = Flask(__name__)


@app.route('/')
def home():
    return render_template('home.html')


@app.route('/docmaker', methods=['POST'])
def pdf_template():
    doctype = request.form.get('doctypes')
    if doctype == 'proofofparticipationinenglishcourses':
        return redirect('/proofofparticipationinenglishcourses')
    elif doctype == 'confirmationofunconditionaladmission':
        return redirect('/confirmationofunconditionaladmission')
    elif doctype == 'extensionletter':
        return redirect('/extensionletter')
    elif doctype == 'latearrival':
        return redirect('/latearrival')
    elif doctype == 'blockedaccount':
        return redirect('/blockedaccount')
    elif doctype == 'onlinesemesterparticipation':
        return redirect('/onlinesemesterparticipation')
    elif doctype == 'presenceletter':
        return redirect('/presenceletter')
    elif doctype == 'proofoflanguagerequirements':
        return redirect('/proofoflanguagerequirements')
    elif doctype == 'virtuallecturesgermany':
        return redirect('/virtuallecturesgermany')
    elif doctype == 'estimatedlivingexpenses':
        return redirect('/estimatedlivingexpenses')
    elif doctype == 'virtuallecturessemester':
        return redirect('/virtuallecturessemester')
    return "Something's wrong"


@app.route('/proofofparticipationinenglishcourses', methods=['GET', 'POST'])
def proofofparticipationinenglishcourses():
    doc = Document(
        './templates/documents/proofofparticipationinenglishcourses.docx')
    if request.method == 'GET':
        return render_template('proofofparticipationinenglishcourses.html')
    else:
        gender = request.form.get('genderattrs')
        name = request.form.get('name') + " " + request.form.get('surname')
        date = time.strftime("%d %B %Y")
        years = request.form.get('yearsattribute')
        studentID = request.form.get('idstudentattribute')
        semester = request.form.get('semesters')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'nameattribute' in text:
                    text = text.replace('nameattribute', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'semesterattribute' in text:
                    text = text.replace('semesterattribute', semester.capitalize())
                    inline[i].text = text
                if 'yearsattribute' in text:
                    text = text.replace('yearsattribute', years)
                    inline[i].text = text
                if 'genderattribute' in text:
                    text = text.replace('genderattribute', gender.capitalize())
                    inline[i].text = text
                if 'idstudentattribute' in text:
                    text = text.replace('idstudentattribute', studentID)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')


@app.route('/confirmationofunconditionaladmission', methods=['GET', 'POST'])
def confirmationofunconditionaladmission():
    doc = Document(
        './templates/documents/confirmationofunconditionaladmission.docx')
    if request.method == 'GET':
        return render_template('confirmationofunconditionaladmission.html')
    else:
        gender = request.form.get('genderattrs')
        name = request.form.get('name') + " " + request.form.get('surname')
        country = request.form.get('countryattribute').upper()
        city = request.form.get('cityattribute')
        zipcode = request.form.get('zipcode')
        address1 = request.form.get('address1')
        address2 = request.form.get('address2')
        number_address = request.form.get('numberaddress')
        date = time.strftime("%d %B %Y")
        years = request.form.get('yearsattribute')
        semester = request.form.get('semesters')
        degreetype = request.form.get('degreetypes')
        degreefield = request.form.get('degreefield')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'nameattr' in text:
                    text = text.replace('nameattr', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'countryattribute' in text:
                    text = text.replace('countryattribute', country)
                    inline[i].text = text
                if 'zipcode' in text:
                    text = text.replace('zipcode', zipcode)
                    inline[i].text = text
                if 'cityattribute' in text:
                    text = text.replace('cityattribute', city)
                    inline[i].text = text
                if 'numberaddress' in text:
                    text = text.replace('numberaddress', number_address)
                    inline[i].text = text
                if 'address1attribute' in text:
                    text = text.replace('address1attribute', address1)
                    inline[i].text = text
                if 'address2attribute' in text:
                    text = text.replace('address2attribute', address2)
                    inline[i].text = text
                if 'semester' in text:
                    text = text.replace('semester', semester.capitalize())
                    inline[i].text = text
                if 'yearsattribute' in text:
                    text = text.replace('yearsattribute', years)
                    inline[i].text = text
                if 'genderattr' in text:
                    text = text.replace('genderattr', gender.capitalize())
                    inline[i].text = text
                if 'degreefield' in text:
                    text = text.replace(
                        'degreefield', degreefield.capitalize())
                    inline[i].text = text
                if 'type' in text:
                    if degreetype == 'bsc':
                        degreetype = 'B.Sc.'
                    else:
                        degreetype = 'M.Sc.'
                    text = text.replace('type', degreetype)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')


@app.route('/extensionletter', methods=['GET', 'POST'])
def extensionletter():
    doc = Document('./templates/documents/extensionletter.docx')
    if request.method == 'GET':
        return render_template('extensionletter.html')
    else:
        name = request.form.get('name') + " " + request.form.get('surname')
        date = time.strftime("%d %B %Y")
        years = request.form.get('yearsattr')
        semester = request.form.get('semesters')
        applicationnum = request.form.get('appnum')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'nametag' in text:
                    text = text.replace('nametag', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'semestera' in text:
                    text = text.replace('semestera', semester.capitalize())
                    inline[i].text = text
                if 'appnum' in text:
                    text = text.replace('appnum', applicationnum)
                    inline[i].text = text
                if 'yearsattr' in text:
                    text = text.replace('yearsattr', years)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')


@app.route('/latearrival', methods=['GET', 'POST'])
def latearrival():
    doc = Document('./templates/documents/latearrival.docx')
    if request.method == 'GET':
        return render_template('latearrival.html')
    else:
        gender = request.form.get('genderattrs')
        name = request.form.get('name') + " " + request.form.get('surname')
        country = request.form.get('countryattribute').upper()
        city = request.form.get('cityattribute')
        zipcode = request.form.get('zipcode')
        address1 = request.form.get('address1')
        address2 = request.form.get('address2')
        number_address = request.form.get('numberaddress')
        date = time.strftime("%d %B %Y")
        studentid = request.form.get('studentid')
        degreefield = request.form.get('degreefield')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'nameattr' in text:
                    text = text.replace('nameattr', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'countryattribute' in text:
                    text = text.replace('countryattribute', country)
                    inline[i].text = text
                if 'zipcode' in text:
                    text = text.replace('zipcode', zipcode)
                    inline[i].text = text
                if 'cityattribute' in text:
                    text = text.replace('cityattribute', city)
                    inline[i].text = text
                if 'numberaddress' in text:
                    text = text.replace('numberaddress', number_address)
                    inline[i].text = text
                if 'address1attribute' in text:
                    text = text.replace('address1attribute', address1)
                    inline[i].text = text
                if 'address2attribute' in text:
                    text = text.replace('address2attribute', address2)
                    inline[i].text = text
                if 'genderattr' in text:
                    text = text.replace('genderattr', gender.capitalize())
                    inline[i].text = text
                if 'studentid' in text:
                    text = text.replace(
                        'studentid', studentid)
                    inline[i].text = text
                if 'degreefield' in text:
                    text = text.replace(
                        'degreefield', degreefield.capitalize())
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')


@app.route('/blockedaccount', methods=['GET', 'POST'])
def blockedaccount():
    doc = Document('./templates/documents/blockedaccount.docx')
    if request.method == 'GET':
        return render_template('blockedaccount.html')
    else:
        date = time.strftime("%d %B %Y")
        years = request.form.get('yearsattr')
        semester = request.form.get('semesters')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'semestera' in text:
                    if semester == 'winter':
                        startingdate = '1st October'
                    else:
                        startingdate = '1st April'
                    text = text.replace('semestera', semester.capitalize())
                    inline[i].text = text
                if 'startingdate' in text:
                    text = text.replace('startingdate', startingdate)
                    inline[i].text = text
                if 'yearsattr' in text:
                    text = text.replace('yearsattr', years)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')


@app.route('/onlinesemesterparticipation', methods=['GET', 'POST'])
def onlinesemesterparticipation():
    doc = Document('./templates/documents/onlinesemesterparticipation.docx')
    if request.method == 'GET':
        return render_template('onlinesemesterparticipation.html')
    else:
        gender = request.form.get('genderattrs')
        name = request.form.get('name') + " " + request.form.get('surname')
        matricnumb = request.form.get('matricnumb')
        numonlinecourse = request.form.get('numonlinecourse')
        date = time.strftime("%d %B %Y")
        semester = request.form.get('semesters')
        year = request.form.get('yearsattribute')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'nameattr' in text:
                    text = text.replace('nameattr', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'matricnum' in text:
                    text = text.replace('matricnum', matricnumb)
                    inline[i].text = text
                if 'semesterattr' in text:
                    text = text.replace('semesterattr', semester.capitalize())
                    inline[i].text = text
                if 'genderattr' in text:
                    text = text.replace('genderattr', gender.capitalize())
                    inline[i].text = text
                if 'numonlinecourse' in text:
                    text = text.replace(
                        'numonlinecourse', numonlinecourse)
                    inline[i].text = text
                if 'yearatt' in text:
                    text = text.replace(
                        'yearatt', year)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')


@app.route('/presenceletter', methods=['GET', 'POST'])
def presenceletter():
    doc = Document('./templates/documents/presenceletter.docx')
    if request.method == 'GET':
        return render_template('presenceletter.html')
    else:
        name = request.form.get('name') + " " + request.form.get('surname')
        date = time.strftime("%d %B %Y")
        years = request.form.get('yearsattr')
        semester = request.form.get('semesters')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'nameattr' in text:
                    text = text.replace('nameattr', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'semesterattribute' in text:
                    text = text.replace('semesterattribute', semester.capitalize())
                    inline[i].text = text
                if 'yearsattr' in text:
                    text = text.replace('yearsattr', years)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')


@app.route('/proofoflanguagerequirements', methods=['GET', 'POST'])
def proofoflanguagerequirements():
    doc = Document('./templates/documents/presenceletter.docx')
    if request.method == 'GET':
        return render_template('proofoflanguagerequirements.html')
    else:
        name = request.form.get('name') + " " + request.form.get('surname')
        date = time.strftime("%d %B %Y")
        years = request.form.get('yearsattr')
        matnum = request.form.get('matnum')
        degreetype = request.form.get('degreetypes')
        degreefield = request.form.get('degreefield')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'nameattr' in text:
                    text = text.replace('nameattr', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'yearsattr' in text:
                    text = text.replace('yearsattr', years)
                    inline[i].text = text
                if 'matnum' in text:
                    text = text.replace('matnum', matnum)
                    inline[i].text = text
                if 'degreefield' in text:
                    text = text.replace(
                        'degreefield', degreefield.capitalize())
                    inline[i].text = text
                if 'tipolaurea' in text:
                    if degreetype == 'bsc':
                        degreetype = 'B.Sc.'
                    else:
                        degreetype = 'M.Sc.'
                    text = text.replace('tipolaurea', degreetype)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')

# #Proof of attendance of solely virtual lectures in a semester
@app.route('/virtuallecturessemester', methods=['GET', 'POST'])
def virtuallecturessemester():
    doc = Document('./templates/documents/virtuallecturessemester.docx')
    if request.method == 'GET':
        return render_template('virtuallecturessemester.html')
    else:
        gender = request.form.get('genderattrs')
        name = request.form.get('name') + " " + request.form.get('surname')
        date = time.strftime("%d %B %Y")
        years = request.form.get('yearsattribute')
        semester = request.form.get('semesters')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'genderattr' in text:
                    text = text.replace('genderattr', gender.capitalize())
                    inline[i].text = text
                if 'nameattr' in text:
                    text = text.replace('nameattr', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'semesterattr' in text:
                    text = text.replace('semesterattr', semester.capitalize())
                    inline[i].text = text
                if 'yearattr' in text:
                    text = text.replace('yearattr', years)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')

# #Proof of attendance of solely virtual lectures until entry to germany
@app.route('/virtuallecturesgermany', methods=['GET', 'POST'])
def virtuallectures():
    doc = Document('./templates/documents/virtuallecturesgermany.docx')
    if request.method == 'GET':
        return render_template('virtuallecturesgermany.html')
    else:
        gender = request.form.get('genderattrs')
        name = request.form.get('name') + " " + request.form.get('surname')
        date = time.strftime("%d %B %Y")
        years = request.form.get('yearsattr')
        semester = request.form.get('semesters')
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'genderattr' in text:
                    text = text.replace('genderattr', gender.capitalize())
                    inline[i].text = text
                if 'nameattr' in text:
                    text = text.replace('nameattr', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
                if 'semesterattr' in text:
                    text = text.replace('semesterattr', semester.capitalize())
                    inline[i].text = text
                if 'yearattr' in text:
                    text = text.replace('yearattr', years)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')

# #Estimated Living Expenses
@app.route('/estimatedlivingexpenses', methods=['GET', 'POST'])
def estimatedlivingexpenses():
    doc = Document('./templates/documents/estimatedlivingexpenses.docx')
    if request.method == 'GET':
        return render_template('estimatedlivingexpenses.html')
    else:
        name = request.form.get('name') + " " + request.form.get('surname')
        date = time.strftime("%d %B %Y")
        for p in doc.paragraphs:
            inline = p.runs
            for i in range(len(inline)):
                text = inline[i].text
                if 'nameattribute' in text:
                    text = text.replace('nameattribute', name)
                    inline[i].text = text
                if 'dateattribute' in text:
                    text = text.replace('dateattribute', date)
                    inline[i].text = text
        doc.save('document.docx')
        return redirect('/download')


@app.route('/download')
def downloadFile():
    return send_file('./document.docx', as_attachment=True)


if __name__ == '__main__':
    app.run()
